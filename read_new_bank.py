import docx
import json
import re

def read_single_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            # 收集后续的描述
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：') or re.match(r'^[A-D]\.', next_text):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            # 收集选项
            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break
            
            # 收集答案
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer = ans_text.replace('答案：', '').strip()
                    i += 1
            
            if question_text and options:
                questions.append({
                    'question': question_text,
                    'options': options,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

def read_multiple_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            # 收集后续的描述
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：') or re.match(r'^[A-D]\.', next_text):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            # 收集选项
            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break
            
            # 收集答案
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer_text = ans_text.replace('答案：', '').strip()
                    # 多选题答案可能是 "AB" 或 ["A","B"]
                    answer = list(answer_text.replace(' ', ''))
                    i += 1
            
            if question_text and options:
                questions.append({
                    'question': question_text,
                    'options': options,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

def read_boolean_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            # 收集后续的描述
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：'):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            # 收集答案
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer = ans_text.replace('答案：', '').strip()
                    i += 1
            
            if question_text:
                questions.append({
                    'question': question_text,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

def read_short_answer_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            # 收集后续的描述
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('> 答案：') or next_text.startswith('> 答案'):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            # 收集答案
            answer = None
            answer_lines = []
            while i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('> 答案：') or ans_text.startswith('> 答案'):
                    answer_line = ans_text.replace('> 答案：', '').replace('> 答案', '').strip()
                    if answer_line:
                        answer_lines.append(answer_line)
                    i += 1
                elif ans_text.startswith('>'):
                    answer_line = ans_text[1:].strip()
                    if answer_line:
                        answer_lines.append(answer_line)
                    i += 1
                else:
                    break
            
            if answer_lines:
                answer = '\n'.join(answer_lines)
            
            if question_text:
                questions.append({
                    'question': question_text,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

print('读取小学单选题...')
primary_single = read_single_questions(r'e:\网页html\青少年人工智能答题\赛道一小学\小学单选题.docx')
print(f'  读取到 {len(primary_single)} 道单选题')

print('读取小学多选题...')
primary_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道一小学\小学多选题.docx')
print(f'  读取到 {len(primary_multiple)} 道多选题')

print('读取小学判断题...')
primary_boolean = read_boolean_questions(r'e:\网页html\青少年人工智能答题\赛道一小学\小学判断题.docx')
print(f'  读取到 {len(primary_boolean)} 道判断题')

print('读取小学主观题...')
primary_short = read_short_answer_questions(r'e:\网页html\青少年人工智能答题\赛道一小学\小学主观题.docx')
print(f'  读取到 {len(primary_short)} 道主观题')

print('\n读取初中单选题...')
junior_single = read_single_questions(r'e:\网页html\青少年人工智能答题\赛道一初中\初中单选题.docx')
print(f'  读取到 {len(junior_single)} 道单选题')

print('读取初中多选题...')
junior_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道一初中\初中多选题.docx')
print(f'  读取到 {len(junior_multiple)} 道多选题')

print('读取初中判断题...')
junior_boolean = read_boolean_questions(r'e:\网页html\青少年人工智能答题\赛道一初中\初中判断题.docx')
print(f'  读取到 {len(junior_boolean)} 道判断题')

print('读取初中主观题...')
junior_short = read_short_answer_questions(r'e:\网页html\青少年人工智能答题\赛道一初中\初中主观题.docx')
print(f'  读取到 {len(junior_short)} 道主观题')
