import docx

doc = docx.Document(r'e:\网页html\青少年人工智能答题\赛道一小学\小学单选题.docx')

# 搜索包含"保护动物"的题目
for i, para in enumerate(doc.paragraphs):
    if '保护动物' in para.text:
        print(f'段落 {i}:')
        print(para.text)
        print()
        
        # 打印前后几个段落看看完整内容
        for j in range(max(0, i-5), min(len(doc.paragraphs), i+10)):
            print(f'  {j}: {doc.paragraphs[j].text}')
        print('-' * 80)
