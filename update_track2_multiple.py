import docx
import json
import re

def read_multiple_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0

    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()

        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1

            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：') or re.match(r'^[A-D]\.', next_text):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1

            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break

            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer_text = ans_text.replace('答案：', '').strip()
                    answer = list(answer_text.replace(' ', ''))
                    i += 1

            if question_text and options:
                questions.append({
                    'question': question_text,
                    'options': options,
                    'answer': answer
                })
        else:
            i += 1

    return questions

# 读取赛道二多选题
print('读取赛道二多选题...')
primary_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道二小学\小学多选题.docx')
junior_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道二初中\初中多选题.docx')

print(f'赛道二小学多选题：{len(primary_multiple)} 道')
print(f'赛道二初中多选题：{len(junior_multiple)} 道')

# 读取现有题库
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 创建题目编号到新题目的映射
primary_map = {}
for q in primary_multiple:
    num_match = re.match(r'^(\d+)', q['question'])
    if num_match:
        num = num_match.group(1)
        primary_map[num] = q

junior_map = {}
for q in junior_multiple:
    num_match = re.match(r'^(\d+)', q['question'])
    if num_match:
        num = num_match.group(1)
        junior_map[num] = q

# 更新赛道二多选题
updated_primary = 0
updated_junior = 0

for q in data:
    if q['type'] != 'multiple':
        continue

    num_match = re.match(r'^(\d+)', q['question'])
    if not num_match:
        continue

    num = num_match.group(1)

    # 更新赛道二小学
    if q['group'] == 'primary' and num in primary_map:
        new_q = primary_map[num]
        q['question'] = new_q['question']
        q['options'] = new_q['options']
        q['answer'] = new_q['answer']
        updated_primary += 1

    # 更新赛道二初中
    elif q['group'] == 'junior' and num in junior_map:
        new_q = junior_map[num]
        q['question'] = new_q['question']
        q['options'] = new_q['options']
        q['answer'] = new_q['answer']
        updated_junior += 1

print(f'\n更新了赛道二小学多选题：{updated_primary} 道')
print(f'更新了赛道二初中多选题：{updated_junior} 道')

# 保存
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'w', encoding='utf-8') as f:
    json.dump(data, f, ensure_ascii=False, indent=2)

print('\n题库已更新！')
