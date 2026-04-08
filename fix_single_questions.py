import docx
import json
import re

# 读取原始题库中的所有单选题
def read_single_questions():
    doc = docx.Document(r'e:\网页html\青少年人工智能答题\赛道一小学\小学单选题.docx')
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        # 检查是否是新题目
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            # 收集后续的描述和选项
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                
                # 如果遇到下一个题目，停止
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：'):
                    break
                
                # 如果是选项，说明描述已经结束
                if re.match(r'^[A-D]\.', next_text):
                    break
                
                # 否则是题目的补充描述
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            # 收集选项
            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break
            
            # 收集答案
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer = ans_text.replace('答案：', '').strip()
                    i += 1
            
            questions.append({
                'question': question_text,
                'options': options,
                'answer': answer
            })
        else:
            i += 1
    
    return questions

# 读取JSON题库
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'r', encoding='utf-8') as f:
    json_data = json.load(f)

# 读取原始单选题
single_questions = read_single_questions()

# 创建题目编号到完整题目的映射
question_map = {}
for q in single_questions:
    num_match = re.match(r'^(\d+)\.', q['question'])
    if num_match:
        num = num_match.group(1)
        question_map[num] = q

# 更新JSON中的题目
updated_count = 0
for q in json_data:
    if q['group'] != 'track1_primary' or q['type'] != 'single':
        continue
    
    num_match = re.match(r'^(\d+)\.', q['question'])
    if not num_match:
        continue
    
    num = num_match.group(1)
    if num not in question_map:
        continue
    
    original = question_map[num]
    
    # 如果原始题目更长，更新
    if len(original['question']) > len(q['question']):
        q['question'] = original['question']
        updated_count += 1
        print(f'更新题目 {num}: {original["question"][:80]}...')

print()
print(f'总共更新了 {updated_count} 道题目')

# 保存更新后的题库
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'w', encoding='utf-8') as f:
    json.dump(json_data, f, ensure_ascii=False, indent=2)

print('题库已更新！')
