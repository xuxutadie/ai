import docx
import json
import re

# 读取多选题
def read_multiple_questions():
    doc = docx.Document(r'e:\网页html\青少年人工智能答题\赛道一小学\小学多选题.docx')
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            # 收集后续的描述
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：') or re.match(r'^[A-D]\.', next_text):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            # 收集选项
            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break
            
            # 收集答案
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer = ans_text.replace('答案：', '').strip()
                    i += 1
            
            questions.append({
                'question': question_text,
                'options': options,
                'answer': answer
            })
        else:
            i += 1
    
    return questions

# 读取判断题
def read_boolean_questions():
    doc = docx.Document(r'e:\网页html\青少年人工智能答题\赛道一小学\小学判断题.docx')
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            # 收集后续的描述
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：'):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            # 收集答案
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer = ans_text.replace('答案：', '').strip()
                    i += 1
            
            questions.append({
                'question': question_text,
                'answer': answer
            })
        else:
            i += 1
    
    return questions

# 读取JSON题库
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'r', encoding='utf-8') as f:
    json_data = json.load(f)

# 更新多选题
multiple_questions = read_multiple_questions()
multiple_map = {}
for q in multiple_questions:
    num_match = re.match(r'^(\d+)\.', q['question'])
    if num_match:
        num = num_match.group(1)
        multiple_map[num] = q

updated_count = 0
for q in json_data:
    if q['group'] != 'track1_primary' or q['type'] != 'multiple':
        continue
    
    num_match = re.match(r'^(\d+)\.', q['question'])
    if not num_match:
        continue
    
    num = num_match.group(1)
    if num not in multiple_map:
        continue
    
    original = multiple_map[num]
    if len(original['question']) > len(q['question']):
        q['question'] = original['question']
        updated_count += 1
        print(f'更新多选题 {num}')

print(f'多选题更新了 {updated_count} 道')

# 更新判断题
boolean_questions = read_boolean_questions()
boolean_map = {}
for q in boolean_questions:
    num_match = re.match(r'^(\d+)\.', q['question'])
    if num_match:
        num = num_match.group(1)
        boolean_map[num] = q

updated_count = 0
for q in json_data:
    if q['group'] != 'track1_primary' or q['type'] != 'boolean':
        continue
    
    num_match = re.match(r'^(\d+)\.', q['question'])
    if not num_match:
        continue
    
    num = num_match.group(1)
    if num not in boolean_map:
        continue
    
    original = boolean_map[num]
    if len(original['question']) > len(q['question']):
        q['question'] = original['question']
        updated_count += 1
        print(f'更新判断题 {num}')

print(f'判断题更新了 {updated_count} 道')

# 保存更新后的题库
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'w', encoding='utf-8') as f:
    json.dump(json_data, f, ensure_ascii=False, indent=2)

print('题库已更新！')
