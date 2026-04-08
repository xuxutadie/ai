import docx
import json
import re

def read_multiple_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0

    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()

        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1

            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：') or re.match(r'^[A-D]\.', next_text):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1

            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break

            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer_text = ans_text.replace('答案：', '').strip()
                    answer = list(answer_text.replace(' ', ''))
                    i += 1

            if question_text and options:
                questions.append({
                    'question': question_text,
                    'options': options,
                    'answer': answer
                })
        else:
            i += 1

    return questions

# 读取赛道二多选题
print('读取赛道二多选题...')
primary_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道二小学\小学多选题.docx')
junior_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道二初中\初中多选题.docx')

print(f'赛道二小学多选题：{len(primary_multiple)} 道')
print(f'赛道二初中多选题：{len(junior_multiple)} 道')

# 读取现有题库
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 移除现有的赛道二多选题
data_no_track2_multiple = [q for q in data if not (q['group'] in ['primary', 'junior'] and q['type'] == 'multiple')]
print(f'\n移除赛道二多选题后剩余：{len(data_no_track2_multiple)} 道')

# 创建新的赛道二多选题
def create_json_questions(read_questions, group, grade, qtype, points):
    json_questions = []
    timestamp = 1775553178000

    for i, q in enumerate(read_questions):
        num_match = re.match(r'^(\d+)', q['question'])
        num = num_match.group(1) if num_match else str(i+1)

        json_q = {
            'id': f'{group}_{qtype}_{timestamp + i}_{num}',
            'grade': grade,
            'group': group,
            'type': qtype,
            'question': q['question'],
            'points': points,
            'answer': q['answer'],
            'options': q['options']
        }

        json_questions.append(json_q)

    return json_questions

# 创建新的赛道二多选题
new_primary_multiple = create_json_questions(primary_multiple, 'primary', '1-6', 'multiple', 2)
new_junior_multiple = create_json_questions(junior_multiple, 'junior', '7-9', 'multiple', 2)

print(f'新建赛道二小学多选题：{len(new_primary_multiple)} 道')
print(f'新建赛道二初中多选题：{len(new_junior_multiple)} 道')

# 合并
new_data = data_no_track2_multiple + new_primary_multiple + new_junior_multiple
print(f'合并后总共：{len(new_data)} 道')

# 保存
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'w', encoding='utf-8') as f:
    json.dump(new_data, f, ensure_ascii=False, indent=2)

print('\n题库已更新！')
