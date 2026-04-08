import docx
import json
import re

def read_single_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：') or re.match(r'^[A-D]\.', next_text):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break
            
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer = ans_text.replace('答案：', '').strip()
                    i += 1
            
            if question_text and options:
                questions.append({
                    'question': question_text,
                    'options': options,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

def read_multiple_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：') or re.match(r'^[A-D]\.', next_text):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break
            
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer_text = ans_text.replace('答案：', '').strip()
                    answer = list(answer_text.replace(' ', ''))
                    i += 1
            
            if question_text and options:
                questions.append({
                    'question': question_text,
                    'options': options,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

def read_boolean_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：'):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer = ans_text.replace('答案：', '').strip()
                    i += 1
            
            if question_text:
                questions.append({
                    'question': question_text,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

def read_fill_in_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：'):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer = ans_text.replace('答案：', '').strip()
                    i += 1
            
            if question_text:
                questions.append({
                    'question': question_text,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

def read_short_answer_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0
    
    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        
        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1
            
            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('> 答案：') or next_text.startswith('> 答案'):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1
            
            answer = None
            answer_lines = []
            while i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('> 答案：') or ans_text.startswith('> 答案'):
                    answer_line = ans_text.replace('> 答案：', '').replace('> 答案', '').strip()
                    if answer_line:
                        answer_lines.append(answer_line)
                    i += 1
                elif ans_text.startswith('>'):
                    answer_line = ans_text[1:].strip()
                    if answer_line:
                        answer_lines.append(answer_line)
                    i += 1
                else:
                    break
            
            if answer_lines:
                answer = '\n'.join(answer_lines)
            
            if question_text:
                questions.append({
                    'question': question_text,
                    'answer': answer
                })
        else:
            i += 1
    
    return questions

def create_json_questions(read_questions, group, grade, qtype, points):
    json_questions = []
    timestamp = 1775553178000
    
    for i, q in enumerate(read_questions):
        num_match = re.match(r'^(\d+)', q['question'])
        num = num_match.group(1) if num_match else str(i+1)
        
        json_q = {
            'id': f'{group}_{qtype}_{timestamp + i}_{num}',
            'grade': grade,
            'group': group,
            'type': qtype,
            'question': q['question'],
            'points': points,
            'answer': q['answer']
        }
        
        if qtype in ['single', 'multiple'] and 'options' in q:
            json_q['options'] = q['options']
        
        json_questions.append(json_q)
    
    return json_questions

# 读取新题库
print('正在读取新题库...')
primary_single = read_single_questions(r'e:\网页html\青少年人工智能答题\赛道一小学新题题库\小学单选题.docx')
primary_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道一小学新题题库\小学多选题.docx')
primary_boolean = read_boolean_questions(r'e:\网页html\青少年人工智能答题\赛道一小学新题题库\小学判断题.docx')
primary_fill = read_fill_in_questions(r'e:\网页html\青少年人工智能答题\赛道一小学新题题库\小学填空题.docx')
primary_short = read_short_answer_questions(r'e:\网页html\青少年人工智能答题\赛道一小学新题题库\小学主观题.docx')

junior_single = read_single_questions(r'e:\网页html\青少年人工智能答题\赛道一初中新题题库\初中单选题.docx')
junior_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道一初中新题题库\初中多选题.docx')
junior_boolean = read_boolean_questions(r'e:\网页html\青少年人工智能答题\赛道一初中新题题库\初中判断题.docx')
junior_fill = read_fill_in_questions(r'e:\网页html\青少年人工智能答题\赛道一初中新题题库\初中填空题.docx')
junior_short = read_short_answer_questions(r'e:\网页html\青少年人工智能答题\赛道一初中新题题库\初中主观题.docx')

print(f'小学：{len(primary_single)}单选、{len(primary_multiple)}多选、{len(primary_boolean)}判断、{len(primary_fill)}填空、{len(primary_short)}主观')
print(f'初中：{len(junior_single)}单选、{len(junior_multiple)}多选、{len(junior_boolean)}判断、{len(junior_fill)}填空、{len(junior_short)}主观')

# 读取现有题库
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

print(f'\n现有题库共 {len(data)} 道题')

# 保留赛道二题目
track2_questions = [q for q in data if q['group'] in ['primary', 'junior']]
print(f'保留赛道二题目：{len(track2_questions)} 道')

# 创建赛道一新题目
print('\n创建赛道一题目...')
track1_questions = []

# 小学题目
track1_questions.extend(create_json_questions(primary_single, 'track1_primary', '1-6', 'single', 2))
track1_questions.extend(create_json_questions(primary_multiple, 'track1_primary', '1-6', 'multiple', 2))
track1_questions.extend(create_json_questions(primary_boolean, 'track1_primary', '1-6', 'boolean', 2))
track1_questions.extend(create_json_questions(primary_fill, 'track1_primary', '1-6', 'fill_in_the_blanks', 2))
track1_questions.extend(create_json_questions(primary_short, 'track1_primary', '1-6', 'short_answer', 10))

# 初中题目
track1_questions.extend(create_json_questions(junior_single, 'track1_junior', '7-9', 'single', 2))
track1_questions.extend(create_json_questions(junior_multiple, 'track1_junior', '7-9', 'multiple', 2))
track1_questions.extend(create_json_questions(junior_boolean, 'track1_junior', '7-9', 'boolean', 2))
track1_questions.extend(create_json_questions(junior_fill, 'track1_junior', '7-9', 'fill_in_the_blanks', 2))
track1_questions.extend(create_json_questions(junior_short, 'track1_junior', '7-9', 'short_answer', 10))

print(f'赛道一题目：{len(track1_questions)} 道')

# 合并
new_data = track2_questions + track1_questions
print(f'\n合并后共 {len(new_data)} 道题')

# 清理无效题目
def is_valid_question(q):
    question = q['question'].strip()
    
    if '文档部分内容可能由 AI 生成' in question:
        return False
    if re.match(r'^[一二三四五六七八九十]+[、.]', question):
        return False
    if question.strip() == '---':
        return False
    if len(question) < 10 and q['type'] not in ['boolean', 'fill_in_the_blanks']:
        has_question_mark = any(c in question for c in ['?', '？'])
        has_content_keywords = any(w in question for w in ['以下', '下列', '关于', '什么', '哪个', '哪些', '是否', '属于', '下列'])
        if not has_question_mark and not has_content_keywords:
            return False
    
    return True

filtered_data = [q for q in new_data if is_valid_question(q)]
print(f'清理后共 {len(filtered_data)} 道题')

# 保存
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'w', encoding='utf-8') as f:
    json.dump(filtered_data, f, ensure_ascii=False, indent=2)

print('\n题库已更新！')
