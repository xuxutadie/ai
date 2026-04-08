import docx
import json
import re

def read_multiple_questions(file_path):
    doc = docx.Document(file_path)
    questions = []
    i = 0

    while i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()

        if re.match(r'^\d+\.', text):
            question_text = text
            i += 1

            while i < len(doc.paragraphs):
                next_text = doc.paragraphs[i].text.strip()
                if re.match(r'^\d+\.', next_text) or next_text.startswith('答案：') or re.match(r'^[A-D]\.', next_text):
                    break
                if next_text:
                    question_text += ' ' + next_text
                i += 1

            options = {}
            while i < len(doc.paragraphs):
                opt_text = doc.paragraphs[i].text.strip()
                opt_match = re.match(r'^([A-D])\.\s*(.*)', opt_text)
                if opt_match:
                    key, val = opt_match.groups()
                    options[key] = val
                    i += 1
                else:
                    break

            answer = None
            if i < len(doc.paragraphs):
                ans_text = doc.paragraphs[i].text.strip()
                if ans_text.startswith('答案：'):
                    answer_text = ans_text.replace('答案：', '').strip()
                    answer = list(answer_text.replace(' ', ''))
                    i += 1

            if question_text and options:
                questions.append({
                    'question': question_text,
                    'options': options,
                    'answer': answer
                })
        else:
            i += 1

    return questions

# 读取赛道二多选题
print('读取赛道二多选题...')
primary_multiple = read_multiple_questions(r'e:\网页html\青少年人工智能答题\赛道二小学\小学多选题.docx')

# 读取现有题库
with open(r'e:\网页html\青少年人工智能答题\src\data\questions.json', 'r', encoding='utf-8') as f:
    data = json.load(f)

# 找到赛道二小学的多选题
existing_primary = [q for q in data if q['group'] == 'primary' and q['type'] == 'multiple']

print(f'\n新题库赛道二小学多选题数量：{len(primary_multiple)}')
print(f'现有题库赛道二小学多选题数量：{len(existing_primary)}')

print('\n新题库前5道题的题号：')
for i, q in enumerate(primary_multiple[:5]):
    num_match = re.match(r'^(\d+)', q['question'])
    num = num_match.group(1) if num_match else 'N/A'
    print(f'  {num}: {q["question"][:60]}...')

print('\n现有题库前5道多选题的题号：')
for i, q in enumerate(existing_primary[:5]):
    num_match = re.match(r'^(\d+)', q['question'])
    num = num_match.group(1) if num_match else 'N/A'
    print(f'  {num}: {q["question"][:60]}...')

# 检查是否有匹配
new_nums = set()
for q in primary_multiple:
    num_match = re.match(r'^(\d+)', q['question'])
    if num_match:
        new_nums.add(num_match.group(1))

existing_nums = set()
for q in existing_primary:
    num_match = re.match(r'^(\d+)', q['question'])
    if num_match:
        existing_nums.add(num_match.group(1))

matched = new_nums & existing_nums
print(f'\n新题库题号数量：{len(new_nums)}')
print(f'现有题库题号数量：{len(existing_nums)}')
print(f'匹配的题号数量：{len(matched)}')
print(f'匹配的题号示例：{list(matched)[:10]}')
